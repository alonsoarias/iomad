#!/usr/bin/env python3
"""
Generador de Informe Técnico - Formato F-GCT-17 ISER
Usa python-docx para editar el documento plantilla existente

Uso:
    python3 generate_report_python.py
"""

import json
import os
from pathlib import Path
from io import BytesIO
from copy import deepcopy

from docx import Document
from docx.shared import Inches, Pt, Cm, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement, register_element_cls
from docx.oxml.text.paragraph import CT_P
from PIL import Image
import cairosvg

# Rutas base
BASE_DIR = Path(__file__).parent
TEMPLATE_PATH = BASE_DIR / 'output' / 'F-GCT-17-Informe-Contratistas-v02-1.docx'
OUTPUT_PATH = BASE_DIR / 'output' / 'INFORME_TECNICO_FGCT17_2025.docx'
DATA_DIR = BASE_DIR / 'data'
SVG_DIR = BASE_DIR / 'svg'
IMAGES_DIR = BASE_DIR / 'images'

# Colores corporativos ISER (RGB)
COLORS = {
    'VERDE': RGBColor(27, 158, 136),
    'AMARILLO': RGBColor(252, 189, 5),
    'ROJO': RGBColor(235, 67, 53),
    'GRIS': RGBColor(100, 99, 99),
    'NEGRO': RGBColor(0, 0, 0),
    'BLANCO': RGBColor(255, 255, 255)
}

# Contadores globales para numeración
figura_counter = 0
tabla_counter = 0

# Registros para las tablas de contenido
figuras_registro = []  # [(numero, titulo), ...]
tablas_registro = []   # [(numero, titulo), ...]


def svg_to_png(svg_path: Path, width: int = 600) -> bytes:
    """Convierte SVG a PNG usando cairosvg."""
    try:
        png_data = cairosvg.svg2png(
            url=str(svg_path),
            output_width=width
        )
        return png_data
    except Exception as e:
        print(f"  ⚠️  Error convirtiendo SVG: {svg_path.name} - {e}")
        return None


def create_element(name):
    """Crea un elemento XML con namespace correcto."""
    return OxmlElement(name)


def create_attribute(element, name, value):
    """Establece un atributo con namespace."""
    element.set(qn(name), value)


def add_field_code(paragraph, field_code: str, display_text: str = ""):
    """
    Agrega un campo de Word (field code) de manera segura.
    Esta implementación crea los elementos XML correctamente.
    """
    run = paragraph.add_run()
    r = run._r

    # fldChar begin
    fld_char_begin = create_element('w:fldChar')
    create_attribute(fld_char_begin, 'w:fldCharType', 'begin')
    r.append(fld_char_begin)

    # Nuevo run para instrText
    run2 = paragraph.add_run()
    r2 = run2._r
    instr_text = create_element('w:instrText')
    create_attribute(instr_text, 'xml:space', 'preserve')
    instr_text.text = f' {field_code} '
    r2.append(instr_text)

    # fldChar separate
    run3 = paragraph.add_run()
    r3 = run3._r
    fld_char_sep = create_element('w:fldChar')
    create_attribute(fld_char_sep, 'w:fldCharType', 'separate')
    r3.append(fld_char_sep)

    # Texto visible
    if display_text:
        run4 = paragraph.add_run(display_text)

    # fldChar end
    run5 = paragraph.add_run()
    r5 = run5._r
    fld_char_end = create_element('w:fldChar')
    create_attribute(fld_char_end, 'w:fldCharType', 'end')
    r5.append(fld_char_end)

    return paragraph


def add_seq_field(paragraph, seq_name: str, number: int):
    """
    Agrega un campo SEQ de Word para numeración automática.
    Cada elemento se agrega en su propio run para evitar corrupción.
    """
    # Run 1: fldChar begin
    r1 = paragraph.add_run()._r
    fld_begin = create_element('w:fldChar')
    create_attribute(fld_begin, 'w:fldCharType', 'begin')
    r1.append(fld_begin)

    # Run 2: instrText
    r2 = paragraph.add_run()._r
    instr = create_element('w:instrText')
    create_attribute(instr, 'xml:space', 'preserve')
    instr.text = f' SEQ {seq_name} \\* ARABIC '
    r2.append(instr)

    # Run 3: fldChar separate
    r3 = paragraph.add_run()._r
    fld_sep = create_element('w:fldChar')
    create_attribute(fld_sep, 'w:fldCharType', 'separate')
    r3.append(fld_sep)

    # Run 4: número visible
    paragraph.add_run(str(number))

    # Run 5: fldChar end
    r5 = paragraph.add_run()._r
    fld_end = create_element('w:fldChar')
    create_attribute(fld_end, 'w:fldCharType', 'end')
    r5.append(fld_end)


def add_caption_with_seq(paragraph, caption_type: str, number: int, title: str):
    """
    Agrega un título con campo SEQ para que Word pueda actualizar las TOC.
    Formato: "Figura SEQ. Título de la figura"
    """
    # Registrar para la tabla de contenido
    if caption_type == 'Figura':
        figuras_registro.append((number, title))
    else:
        tablas_registro.append((number, title))

    # Agregar prefijo "Figura " o "Tabla "
    run_prefix = paragraph.add_run(f'{caption_type} ')
    run_prefix.italic = True
    run_prefix.font.size = Pt(10)
    run_prefix.font.color.rgb = COLORS['GRIS']

    # Agregar campo SEQ
    add_seq_field(paragraph, caption_type, number)

    # Agregar ". título"
    run_suffix = paragraph.add_run(f'. {title}')
    run_suffix.italic = True
    run_suffix.font.size = Pt(10)
    run_suffix.font.color.rgb = COLORS['GRIS']

    return paragraph


def add_simple_caption(paragraph, caption_type: str, number: int, title: str):
    """
    Agrega un título con numeración usando campos SEQ de Word.
    Esto permite que Word actualice automáticamente las tablas de contenido.
    """
    return add_caption_with_seq(paragraph, caption_type, number, title)


def add_figure_caption(doc, title: str):
    """Agrega un título de figura con numeración automática y fuente."""
    global figura_counter
    figura_counter += 1

    # Título de la figura
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_simple_caption(caption, 'Figura', figura_counter, title)

    # Fuente: Elaboración propia
    source = doc.add_paragraph()
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = source.add_run('Fuente: Elaboración propia')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = COLORS['GRIS']
    source.paragraph_format.space_before = Pt(0)
    source.paragraph_format.space_after = Pt(6)

    return caption


def add_table_caption(doc, title: str):
    """Agrega un título de tabla con numeración automática."""
    global tabla_counter
    tabla_counter += 1

    caption = doc.add_paragraph()
    add_simple_caption(caption, 'Tabla', tabla_counter, title)

    return caption


def add_toc_field(doc, toc_type: str):
    """
    Agrega un campo TOC para tabla de contenido.
    toc_type: 'content', 'figure', 'table'
    """
    paragraph = doc.add_paragraph()

    if toc_type == 'content':
        field_code = 'TOC \\o "1-3" \\h \\z \\u'
        display = "Actualice el campo presionando F9"
    elif toc_type == 'figure':
        field_code = 'TOC \\h \\z \\c "Figura"'
        display = "Actualice el campo presionando F9"
    elif toc_type == 'table':
        field_code = 'TOC \\h \\z \\c "Tabla"'
        display = "Actualice el campo presionando F9"
    else:
        return paragraph

    add_field_code(paragraph, field_code, display)

    return paragraph


def add_figure(doc, svg_path: Path, title: str, width_inches: float = 5.5):
    """Agrega una figura con su título."""
    if not svg_path.exists():
        print(f"  ⚠️  SVG no encontrado: {svg_path}")
        return

    # Convertir SVG a PNG
    png_data = svg_to_png(svg_path, width=int(width_inches * 150))
    if png_data is None:
        return

    # Agregar imagen centrada
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(BytesIO(png_data), width=Inches(width_inches))

    # Agregar título
    add_figure_caption(doc, title)

    # Espacio después
    doc.add_paragraph()


def add_table_with_caption(doc, data: list, title: str, col_widths: list = None):
    """Agrega una tabla con su título DEBAJO y formato corporativo."""
    if not data:
        return

    # Crear tabla PRIMERO
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Llenar datos
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_data) if cell_data else ''

            # Aplicar fuente Arial a todas las celdas
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)

            # Estilo de encabezado (primera fila)
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = COLORS['BLANCO']
                # Fondo verde para encabezado
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), '1B9E88')  # Verde ISER
                cell._tc.get_or_add_tcPr().append(shading)

    # Aplicar anchos de columna si se especifican
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                if i < len(row.cells):
                    row.cells[i].width = Inches(width)

    # Título de la tabla DESPUÉS de la tabla
    add_table_caption(doc, title)

    doc.add_paragraph()  # Espacio después del título


def load_json_data(plugin_name: str) -> dict:
    """Carga los datos JSON consolidados de un plugin."""
    json_path = DATA_DIR / plugin_name / 'CONSOLIDADO.json'
    if not json_path.exists():
        print(f"  ⚠️  JSON no encontrado: {json_path}")
        return {}

    with open(json_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def ensure_heading_styles(doc):
    """
    Asegura que los estilos de encabezado existan en el documento.
    Crea los estilos Heading 1, 2, 3 si no existen.
    """
    styles = doc.styles

    heading_configs = {
        'Heading 1': {'size': 16, 'color': COLORS['VERDE'], 'bold': True},
        'Heading 2': {'size': 14, 'color': COLORS['NEGRO'], 'bold': True},
        'Heading 3': {'size': 12, 'color': COLORS['NEGRO'], 'bold': True},
    }

    for style_name, config in heading_configs.items():
        try:
            style = styles[style_name]
        except KeyError:
            # Crear el estilo si no existe
            from docx.enum.style import WD_STYLE_TYPE
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles['Normal']

        # Configurar el estilo
        style.font.bold = config['bold']
        style.font.size = Pt(config['size'])
        style.font.color.rgb = config['color']


def set_paragraph_outline_level(paragraph, level: int):
    """
    Establece el nivel de esquema (outline level) del párrafo.
    Esto es necesario para que Word reconozca el párrafo como encabezado en la TOC.
    """
    pPr = paragraph._p.get_or_add_pPr()

    # Buscar o crear el elemento outlineLvl
    outline_lvl = pPr.find(qn('w:outlineLvl'))
    if outline_lvl is None:
        outline_lvl = create_element('w:outlineLvl')
        pPr.append(outline_lvl)

    # El nivel de outline es 0-indexed (Heading 1 = nivel 0)
    create_attribute(outline_lvl, 'w:val', str(level - 1))


def add_heading(doc, text: str, level: int = 1):
    """
    Agrega un encabezado con estilo de Word para la TOC.
    Usa estilos Heading 1/2/3 y establece outline level.
    Aplica colores institucionales ISER.
    """
    # Crear párrafo y aplicar formato manualmente para control total de colores
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True

    # Tamaños según nivel
    sizes = {1: 16, 2: 14, 3: 12}
    run.font.size = Pt(sizes.get(level, 12))
    run.font.name = 'Arial'

    # Colores institucionales ISER según nivel
    if level == 1:
        run.font.color.rgb = COLORS['VERDE']  # Verde ISER para títulos principales
    elif level == 2:
        run.font.color.rgb = COLORS['GRIS']   # Gris para subtítulos
    else:
        run.font.color.rgb = COLORS['NEGRO']  # Negro para nivel 3

    # Establecer outline level para TOC (crítico para que funcione)
    set_paragraph_outline_level(p, level)

    # Formato de espaciado
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)

    return p


def add_paragraph_text(doc, text: str, bold: bool = False, italic: bool = False):
    """Agrega un párrafo de texto con formato corporativo."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    # Justificado con interlineado 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15

    return p


def add_bullet_list(doc, items: list):
    """Agrega una lista con viñetas con formato corporativo."""
    for item in items:
        p = doc.add_paragraph()
        run = p.add_run(f'• {item}')
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15


def generate_cover_page(doc):
    """Genera la portada institucional del documento."""
    # Encabezado institucional
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run('INSTITUTO SUPERIOR DE EDUCACIÓN RURAL')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Arial'
    run.font.color.rgb = COLORS['VERDE']

    header2 = doc.add_paragraph()
    header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header2.add_run('ISER')
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Arial'
    run.font.color.rgb = COLORS['VERDE']

    # Línea decorativa
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run('═' * 50)
    run.font.size = Pt(12)
    run.font.color.rgb = COLORS['AMARILLO']

    vice = doc.add_paragraph()
    vice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = vice.add_run('VICERRECTORÍA ACADÉMICA')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    run.font.color.rgb = COLORS['GRIS']

    for _ in range(3):
        doc.add_paragraph()

    # Título del documento
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('INFORME TÉCNICO')
    run.bold = True
    run.font.size = Pt(26)
    run.font.name = 'Arial'
    run.font.color.rgb = COLORS['VERDE']

    # Subtítulo
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Documentación Técnica de Plugins Moodle')
    run.font.size = Pt(14)
    run.font.name = 'Arial'

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run('Plataforma Virtual ISER')
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    run.font.color.rgb = COLORS['GRIS']

    for _ in range(2):
        doc.add_paragraph()

    # Recuadro de plugins
    box_line = doc.add_paragraph()
    box_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = box_line.add_run('┌' + '─' * 40 + '┐')
    run.font.size = Pt(10)
    run.font.name = 'Consolas'
    run.font.color.rgb = COLORS['VERDE']

    plugins_header = doc.add_paragraph()
    plugins_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = plugins_header.add_run('│       COMPONENTES DOCUMENTADOS       │')
    run.font.size = Pt(10)
    run.font.name = 'Consolas'
    run.font.color.rgb = COLORS['VERDE']

    plugins = ['local_jobboard', 'report_platform_usage', 'local_platform_access']
    for plugin in plugins:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        padding = ' ' * ((38 - len(plugin)) // 2)
        run = p.add_run(f'│{padding}{plugin}{padding}│')
        run.font.size = Pt(10)
        run.font.name = 'Consolas'
        run.font.color.rgb = COLORS['VERDE']

    box_end = doc.add_paragraph()
    box_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = box_end.add_run('└' + '─' * 40 + '┘')
    run.font.size = Pt(10)
    run.font.name = 'Consolas'
    run.font.color.rgb = COLORS['VERDE']

    for _ in range(3):
        doc.add_paragraph()

    # Marco del contrato
    contract_frame = doc.add_paragraph()
    contract_frame.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contract_frame.add_run('CONTRATO DE PRESTACIÓN DE SERVICIOS')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.font.color.rgb = COLORS['GRIS']

    contract_num = doc.add_paragraph()
    contract_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contract_num.add_run('No. CD-CPS-066 de 2025')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    run.font.color.rgb = COLORS['VERDE']

    for _ in range(2):
        doc.add_paragraph()

    # Información del contratista
    presented = doc.add_paragraph()
    presented.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = presented.add_run('Presentado por:')
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    run.font.color.rgb = COLORS['GRIS']

    author_name = doc.add_paragraph()
    author_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_name.add_run('PEDRO ALONSO ARIAS BALCUCHO')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'

    author_cc = doc.add_paragraph()
    author_cc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_cc.add_run('C.C. 1.094.277.917')
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.font.color.rgb = COLORS['GRIS']

    for _ in range(3):
        doc.add_paragraph()

    # Línea decorativa inferior
    line2 = doc.add_paragraph()
    line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line2.add_run('═' * 50)
    run.font.size = Pt(12)
    run.font.color.rgb = COLORS['AMARILLO']

    # Lugar y fecha
    location = doc.add_paragraph()
    location.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = location.add_run('Pamplona, Norte de Santander')
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.font.color.rgb = COLORS['GRIS']

    year = doc.add_paragraph()
    year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = year.add_run('Diciembre 2025')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    run.font.color.rgb = COLORS['VERDE']

    doc.add_page_break()


def generate_toc_pages(doc):
    """Genera las páginas de tablas de contenido."""
    # Tabla de contenido
    add_heading(doc, 'TABLA DE CONTENIDO', 1)
    add_toc_field(doc, 'content')
    doc.add_page_break()

    # Lista de ilustraciones
    add_heading(doc, 'LISTA DE ILUSTRACIONES', 1)
    add_toc_field(doc, 'figure')
    doc.add_page_break()

    # Lista de tablas
    add_heading(doc, 'LISTA DE TABLAS', 1)
    add_toc_field(doc, 'table')
    doc.add_page_break()


def generate_introduction(doc):
    """Genera las secciones introductorias."""
    add_heading(doc, '1. INTRODUCCIÓN', 1)

    add_paragraph_text(doc,
        'El presente documento técnico describe la implementación, arquitectura y '
        'funcionamiento de tres plugins desarrollados para la plataforma Moodle del '
        'Instituto Superior de Educación Rural (ISER). Estos plugins extienden las '
        'funcionalidades de la plataforma virtual para satisfacer necesidades específicas '
        'de la institución en áreas como la gestión de vacantes docentes, reportes de uso '
        'y generación de datos de acceso para pruebas.')

    add_heading(doc, '2. OBJETIVOS', 1)

    add_heading(doc, '2.1 Objetivo General', 2)
    add_paragraph_text(doc,
        'Documentar técnicamente los plugins desarrollados para la plataforma Moodle ISER, '
        'proporcionando información detallada sobre su arquitectura, funcionamiento e '
        'integración con el sistema, de manera que sirva como referencia para su mantenimiento '
        'y evolución futura.')

    add_heading(doc, '2.2 Objetivos Específicos', 2)
    objectives = [
        'Describir la estructura y organización de archivos de cada plugin',
        'Documentar la arquitectura de base de datos utilizada',
        'Explicar los flujos de trabajo y casos de uso implementados',
        'Detallar las APIs y servicios web disponibles',
        'Proporcionar guías de mantenimiento y extensión del código'
    ]
    add_bullet_list(doc, objectives)

    add_heading(doc, '3. ALCANCE', 1)
    add_paragraph_text(doc,
        'Este documento cubre la documentación técnica completa de los siguientes plugins '
        'desarrollados específicamente para el ISER:')

    plugins_info = [
        'local_jobboard: Sistema de gestión de vacantes y postulaciones docentes',
        'report_platform_usage: Generador de reportes de uso de la plataforma',
        'local_platform_access: Generador de registros de acceso para pruebas de carga'
    ]
    add_bullet_list(doc, plugins_info)

    add_heading(doc, '4. TECNOLOGÍAS UTILIZADAS', 1)
    add_paragraph_text(doc,
        'Los plugins han sido desarrollados siguiendo los estándares de desarrollo de Moodle '
        'y utilizando las siguientes tecnologías:')

    tech_data = [
        ['Categoría', 'Tecnología', 'Descripción'],
        ['Backend', 'PHP 8.x', 'Lenguaje principal del core de Moodle'],
        ['Base de datos', 'MySQL/MariaDB', 'Acceso mediante DBAL de Moodle'],
        ['Frontend', 'JavaScript (AMD)', 'Módulos RequireJS para interactividad'],
        ['Plantillas', 'Mustache', 'Motor de plantillas nativo de Moodle'],
        ['Web Services', 'REST/AJAX', 'API External de Moodle'],
        ['Caché', 'MUC', 'Moodle Universal Cache para optimización']
    ]
    add_table_with_caption(doc, tech_data, 'Tecnologías utilizadas en los plugins')

    doc.add_page_break()


def generate_jobboard_section(doc):
    """Genera la sección del plugin local_jobboard."""
    data = load_json_data('local_jobboard')
    if not data:
        return

    add_heading(doc, '5. PLUGIN LOCAL_JOBBOARD', 1)

    # Descripción
    add_heading(doc, '5.1 Descripción General', 2)
    if 'version' in data:
        desc = data['version'].get('description', '')
        if desc:
            add_paragraph_text(doc, desc)

    add_paragraph_text(doc,
        'Este plugin proporciona un sistema completo para la publicación de vacantes docentes, '
        'gestión de postulaciones, revisión de documentos y seguimiento de todo el proceso de '
        'selección de personal académico.')

    # Información de versión
    add_heading(doc, '5.2 Información de Versión', 2)
    if 'version' in data:
        v = data['version']
        version_data = [
            ['Campo', 'Valor'],
            ['Componente', 'local_jobboard'],
            ['Versión', v.get('release', 'N/A')],
            ['Build', str(v.get('version', 'N/A'))],
            ['Madurez', v.get('maturity', 'N/A')],
            ['Moodle requerido', v.get('requires_description', 'N/A')],
            ['Autor', v.get('author', 'N/A')],
            ['Licencia', v.get('license', 'GPLv3')]
        ]
        add_table_with_caption(doc, version_data, 'Información de versión del plugin local_jobboard')

    # Estructura del plugin
    add_heading(doc, '5.3 Estructura del Plugin', 2)
    add_paragraph_text(doc,
        'El plugin local_jobboard sigue la estructura estándar de plugins locales de Moodle. '
        'La organización de directorios y archivos está diseñada para facilitar el mantenimiento '
        'y cumplir con las convenciones de desarrollo de Moodle.')

    svg_path = SVG_DIR / 'local_jobboard' / 'estructura_directorios.svg'
    add_figure(doc, svg_path, 'Estructura de directorios del plugin local_jobboard')

    # Arquitectura
    add_heading(doc, '5.4 Arquitectura del Sistema', 2)
    add_paragraph_text(doc,
        'El plugin implementa una arquitectura modular que separa las responsabilidades en capas: '
        'presentación (templates Mustache y JavaScript), lógica de negocio (clases PHP), '
        'acceso a datos (Moodle DBAL), y servicios web (External API).')

    svg_path = SVG_DIR / 'local_jobboard' / 'arquitectura.svg'
    add_figure(doc, svg_path, 'Arquitectura del plugin local_jobboard')

    # Casos de uso
    add_heading(doc, '5.5 Casos de Uso', 2)
    add_paragraph_text(doc,
        'El sistema define cinco tipos de actores principales: Docente/Aspirante (quien postula), '
        'Revisor de Documentos (valida requisitos), Administrador (gestiona el sistema), '
        'Decano de Facultad (aprueba contrataciones) y Recursos Humanos (formaliza contratos).')

    svg_path = SVG_DIR / 'local_jobboard' / 'casos_uso.svg'
    add_figure(doc, svg_path, 'Diagrama de casos de uso del sistema de vacantes')

    # Base de datos
    add_heading(doc, '5.6 Base de Datos', 2)
    if 'database' in data:
        tables = data['database'].get('tables', [])
        add_paragraph_text(doc,
            f'El plugin utiliza {len(tables)} tablas para almacenar la información de vacantes, '
            'postulaciones, documentos y configuraciones del sistema.')

        table_summary = [['Tabla', 'Descripción']]
        for t in tables[:8]:
            name = t.get('name', '').replace('local_jobboard_', '')
            comment = t.get('comment', 'Sin descripción')
            table_summary.append([name, comment[:60] + '...' if len(comment) > 60 else comment])

        add_table_with_caption(doc, table_summary, 'Resumen de tablas de base de datos del plugin local_jobboard')

    # Diagrama ER
    svg_path = SVG_DIR / 'local_jobboard' / 'diagrama_er.svg'
    add_figure(doc, svg_path, 'Diagrama entidad-relación del plugin local_jobboard')

    # Diagrama de clases
    add_heading(doc, '5.7 Clases Principales', 2)
    add_paragraph_text(doc,
        'El plugin define varias clases PHP que encapsulan la lógica de negocio, '
        'siguiendo los patrones de diseño recomendados por Moodle.')

    svg_path = SVG_DIR / 'local_jobboard' / 'diagrama_clases.svg'
    add_figure(doc, svg_path, 'Diagrama de clases principales del plugin local_jobboard')

    # Servicios web
    add_heading(doc, '5.8 Servicios Web', 2)
    if 'services' in data and data['services']:
        add_paragraph_text(doc,
            f'El plugin expone {len(data["services"])} servicios web que permiten la interacción '
            'con el sistema mediante llamadas AJAX desde el frontend y APIs externas.')

        svg_path = SVG_DIR / 'local_jobboard' / 'servicios_web.svg'
        add_figure(doc, svg_path, 'Servicios web disponibles del plugin local_jobboard')

        services_data = [['Función', 'Tipo', 'Descripción']]
        for s in data['services'][:6]:
            services_data.append([
                s.get('function_name', '')[:30],
                s.get('type', ''),
                s.get('description', '')[:40] + '...' if len(s.get('description', '')) > 40 else s.get('description', '')
            ])
        add_table_with_caption(doc, services_data, 'Servicios web del plugin local_jobboard')

    # Flujos de trabajo
    add_heading(doc, '5.9 Flujos de Trabajo', 2)

    add_heading(doc, '5.9.1 Flujo de Estados de Vacante', 3)
    add_paragraph_text(doc,
        'Las vacantes siguen un ciclo de vida definido que va desde su creación como borrador '
        'hasta su cierre una vez completado el proceso de selección.')

    svg_path = SVG_DIR / 'local_jobboard' / 'flujo_estados_vacante.svg'
    add_figure(doc, svg_path, 'Flujo de estados de una vacante')

    add_heading(doc, '5.9.2 Flujo de Postulación', 3)
    add_paragraph_text(doc,
        'El proceso de postulación permite a los docentes enviar su candidatura, '
        'adjuntar documentos requeridos y hacer seguimiento del estado de su aplicación.')

    svg_path = SVG_DIR / 'local_jobboard' / 'flujo_postulacion.svg'
    add_figure(doc, svg_path, 'Flujo del proceso de postulación')

    add_heading(doc, '5.9.3 Ciclo de Vida de la Aplicación', 3)
    svg_path = SVG_DIR / 'local_jobboard' / 'ciclo_vida_aplicacion.svg'
    add_figure(doc, svg_path, 'Ciclo de vida completo de una aplicación')

    # Eventos del sistema
    add_heading(doc, '5.10 Eventos del Sistema', 2)
    if 'events' in data:
        events = data['events']
        add_paragraph_text(doc,
            'El plugin dispara eventos del sistema de Moodle para permitir la integración '
            'con otros componentes y el registro de actividad en los logs.')

        if 'eventos' in events:
            events_data = [['Evento', 'Tipo', 'Descripción']]
            for ev in events['eventos'][:8]:
                events_data.append([
                    ev.get('nombre', 'N/A'),
                    ev.get('tipo_accion', 'N/A'),
                    ev.get('descripcion', 'N/A')[:40] + '...' if len(ev.get('descripcion', '')) > 40 else ev.get('descripcion', 'N/A')
                ])
            add_table_with_caption(doc, events_data, 'Eventos del sistema del plugin local_jobboard')

    # Templates de notificación
    add_heading(doc, '5.11 Sistema de Notificaciones', 2)
    add_paragraph_text(doc,
        'El plugin incluye un sistema de plantillas de correo electrónico personalizables '
        'con soporte multi-tenant para IOMAD. Las plantillas permiten usar marcadores '
        'que se reemplazan dinámicamente con información del contexto.')

    notification_cats = [
        'application: Notificaciones relacionadas con postulaciones',
        'documents: Notificaciones sobre documentos y validación',
        'interview: Notificaciones de programación de entrevistas',
        'selection: Notificaciones de resultados de selección',
        'system: Notificaciones del sistema'
    ]
    add_bullet_list(doc, notification_cats)

    # Capabilities
    add_heading(doc, '5.12 Sistema de Permisos', 2)
    if 'capabilities' in data:
        caps_list = data['capabilities']
        # Contar todas las capabilities en las categorías
        total_caps = sum(len(cat.get('capabilities', [])) for cat in caps_list if isinstance(cat, dict))
        add_paragraph_text(doc,
            f'El plugin define {total_caps} capabilities organizadas en {len(caps_list)} categorías '
            'para controlar el acceso a las diferentes funcionalidades del sistema.')

        caps_data = [['Categoría', 'Capability', 'Tipo']]
        for cat in caps_list[:6]:
            cat_name = cat.get('name', 'N/A')
            for cap in cat.get('capabilities', [])[:2]:
                caps_data.append([
                    cat_name,
                    cap.get('name', 'N/A').replace('local/jobboard:', ''),
                    cap.get('type', 'N/A')
                ])
        add_table_with_caption(doc, caps_data, 'Principales capabilities del plugin local_jobboard')

    # Templates Mustache
    add_heading(doc, '5.13 Plantillas Mustache', 2)
    if 'templates' in data:
        templates = data['templates']
        total = templates.get('metadata', {}).get('total_templates', 0)
        add_paragraph_text(doc,
            f'El plugin utiliza {total} plantillas Mustache organizadas por categoría para '
            'la capa de presentación, siguiendo el patrón de separación de responsabilidades de Moodle.')

        if 'template_categories' in templates:
            tpl_data = [['Categoría', 'Descripción', 'Cantidad']]
            for cat in templates['template_categories']:
                tpl_data.append([
                    cat.get('category', 'N/A'),
                    cat.get('description', 'N/A')[:40] + '...' if len(cat.get('description', '')) > 40 else cat.get('description', 'N/A'),
                    str(len(cat.get('templates', [])))
                ])
            add_table_with_caption(doc, tpl_data, 'Categorías de templates Mustache del plugin local_jobboard')

    # Configuraciones del plugin
    add_heading(doc, '5.14 Configuraciones del Plugin', 2)
    if 'settings' in data:
        settings = data['settings']
        summary = settings.get('summary', {})
        add_paragraph_text(doc,
            f'El plugin ofrece {summary.get("total_settings", 0)} configuraciones organizadas en '
            f'{summary.get("total_sections", 0)} secciones, además de '
            f'{summary.get("total_external_pages", 0)} páginas de administración externas.')

        if 'settings_sections' in settings:
            settings_data = [['Sección', 'Configuraciones', 'Propósito']]
            for section in settings['settings_sections'][:8]:
                section_name = section.get('section', 'N/A')
                setting_list = section.get('settings', [])
                num_settings = len(setting_list)
                purpose = ''
                if setting_list:
                    purpose = setting_list[0].get('purpose', 'N/A')[:35] + '...' if len(setting_list[0].get('purpose', '')) > 35 else setting_list[0].get('purpose', 'N/A')
                settings_data.append([section_name, str(num_settings), purpose])
            add_table_with_caption(doc, settings_data, 'Secciones de configuración del plugin local_jobboard')

        if 'external_pages' in settings:
            ext_pages_data = [['Página', 'URL', 'Descripción']]
            for page in settings['external_pages']:
                ext_pages_data.append([
                    page.get('name', 'N/A').replace('local_jobboard_', ''),
                    page.get('url', 'N/A'),
                    page.get('description', 'N/A')[:40] + '...' if len(page.get('description', '')) > 40 else page.get('description', 'N/A')
                ])
            add_table_with_caption(doc, ext_pages_data, 'Páginas de administración externas del plugin local_jobboard')

    # Resumen del plugin
    add_heading(doc, '5.15 Resumen Técnico', 2)
    if 'summary' in data:
        summary = data['summary']
        add_paragraph_text(doc,
            'El siguiente resumen proporciona una visión general de las métricas técnicas '
            'del plugin local_jobboard:')

        # El summary tiene total_tables, total_capabilities, etc.
        summary_list = []
        if 'total_tables' in summary:
            summary_list.append(f'Total de tablas en base de datos: {summary["total_tables"]}')
        if 'total_capabilities' in summary:
            summary_list.append(f'Capabilities definidas: {summary["total_capabilities"]}')
        if 'total_services' in summary:
            summary_list.append(f'Servicios web: {summary["total_services"]}')
        if 'total_functions' in summary:
            summary_list.append(f'Funciones PHP: {summary["total_functions"]}')
        if summary_list:
            add_bullet_list(doc, summary_list)

    # Clases principales
    add_heading(doc, '5.16 Clases Principales', 2)
    if 'classes' in data:
        classes = data['classes']
        add_paragraph_text(doc,
            'El plugin implementa las siguientes clases PHP principales:')

        classes_data = [['Clase', 'Namespace', 'Descripción']]
        for cls_name, cls_info in classes.items():
            if isinstance(cls_info, dict):
                classes_data.append([
                    cls_info.get('clase', cls_name),
                    cls_info.get('namespace', 'local_jobboard'),
                    cls_info.get('descripcion', 'N/A')[:50] + '...' if len(cls_info.get('descripcion', '')) > 50 else cls_info.get('descripcion', 'N/A')
                ])
        if len(classes_data) > 1:
            add_table_with_caption(doc, classes_data, 'Clases principales del plugin local_jobboard')

    doc.add_page_break()


def generate_platform_usage_section(doc):
    """Genera la sección del plugin report_platform_usage."""
    data = load_json_data('report_platform_usage')
    if not data:
        return

    add_heading(doc, '6. PLUGIN REPORT_PLATFORM_USAGE', 1)

    # Introducción
    add_heading(doc, '6.1 Descripción General', 2)
    if 'resumen_ejecutivo' in data:
        resumen = data['resumen_ejecutivo']
        add_paragraph_text(doc,
            f'{resumen.get("nombre_plugin", "Report Platform Usage")} es un plugin de tipo '
            f'{resumen.get("tipo", "reporte")} diseñado para proporcionar '
            f'{resumen.get("proposito", "informes detallados sobre el uso de la plataforma")}.')

    add_paragraph_text(doc,
        'Este plugin permite a los administradores generar reportes estadísticos sobre '
        'el uso de la plataforma, incluyendo métricas de acceso, actividad de usuarios '
        'y utilización de recursos educativos.')

    # Estructura
    add_heading(doc, '6.2 Estructura del Plugin', 2)
    svg_path = SVG_DIR / 'report_platform_usage' / 'estructura_directorios.svg'
    add_figure(doc, svg_path, 'Estructura de directorios del plugin report_platform_usage')

    # Información de versión
    add_heading(doc, '6.3 Información de Versión', 2)
    if 'informacion_version' in data:
        v = data['informacion_version']
        version_data = [
            ['Atributo', 'Valor'],
            ['Plugin', v.get('plugin', 'report_platform_usage')],
            ['Versión', v.get('version_release', 'N/A')],
            ['Moodle requerido', v.get('moodle_requerido', 'N/A')],
            ['Madurez', v.get('madurez', 'N/A')],
            ['Licencia', v.get('licencia', 'GPLv3')]
        ]
        add_table_with_caption(doc, version_data, 'Información de versión del plugin report_platform_usage')

    # Arquitectura
    add_heading(doc, '6.4 Arquitectura del Sistema', 2)
    add_paragraph_text(doc,
        'El plugin sigue la arquitectura estándar de reportes de Moodle, integrándose '
        'con el sistema de reportes del sitio y utilizando las APIs de logging para '
        'extraer información de uso.')

    svg_path = SVG_DIR / 'report_platform_usage' / 'arquitectura.svg'
    add_figure(doc, svg_path, 'Diagrama de arquitectura del plugin report_platform_usage')

    # Clases principales
    add_heading(doc, '6.5 Clases Principales', 2)
    if 'arquitectura_clases' in data and 'clases_principales' in data['arquitectura_clases']:
        clases = data['arquitectura_clases']['clases_principales']
        add_paragraph_text(doc,
            'El plugin implementa varias clases PHP especializadas que encapsulan '
            'la lógica de negocio y el acceso a datos.')

        clases_data = [['Clase', 'Responsabilidad']]
        for cls in clases:
            clases_data.append([
                cls.get('nombre', 'N/A'),
                cls.get('responsabilidad', 'N/A')[:50] + '...' if len(cls.get('responsabilidad', '')) > 50 else cls.get('responsabilidad', 'N/A')
            ])
        add_table_with_caption(doc, clases_data, 'Clases principales del plugin report_platform_usage')

    # Base de datos
    add_heading(doc, '6.6 Arquitectura de Base de Datos', 2)
    if 'arquitectura_base_datos' in data:
        db = data['arquitectura_base_datos']
        if 'tablas_propias' in db:
            for tabla in db['tablas_propias']:
                tabla_data = [
                    ['Campo', 'Descripción'],
                    ['Nombre', tabla.get('nombre', 'N/A')],
                    ['Propósito', tabla.get('proposito', 'N/A')],
                    ['Poblada por', tabla.get('poblada_por', 'N/A')],
                    ['Frecuencia', tabla.get('frecuencia_actualizacion', 'N/A')]
                ]
                add_table_with_caption(doc, tabla_data, f'Estructura de la tabla {tabla.get("nombre", "de datos")}')

    # Capabilities
    add_heading(doc, '6.7 Capabilities y Seguridad', 2)
    add_paragraph_text(doc,
        'El plugin define capabilities específicas para controlar el acceso a los reportes '
        'y la visualización de información sensible.')

    if 'seguridad_permisos' in data and 'capabilities' in data['seguridad_permisos']:
        caps = data['seguridad_permisos']['capabilities']
        caps_data = [['Capability', 'Tipo', 'Propósito']]
        for cap in caps:
            caps_data.append([
                cap.get('nombre', 'N/A'),
                cap.get('tipo', 'N/A'),
                cap.get('proposito', 'N/A')
            ])
        add_table_with_caption(doc, caps_data, 'Capabilities del plugin report_platform_usage')

    # Flujo de datos
    add_heading(doc, '6.8 Flujo de Datos', 2)
    add_paragraph_text(doc,
        'El reporte recopila datos de múltiples tablas del sistema de logging de Moodle, '
        'los procesa y presenta en formatos tabulares y gráficos.')

    svg_path = SVG_DIR / 'report_platform_usage' / 'flujo_datos.svg'
    add_figure(doc, svg_path, 'Diagrama de flujo de datos del plugin report_platform_usage')

    # Flujo del reporte
    add_heading(doc, '6.9 Flujo del Reporte', 2)
    add_paragraph_text(doc,
        'El flujo de generación de reportes incluye la recopilación de métricas, '
        'el procesamiento de datos y la presentación en múltiples formatos de exportación.')

    svg_path = SVG_DIR / 'report_platform_usage' / 'flujo_reporte.svg'
    add_figure(doc, svg_path, 'Diagrama de flujo del reporte de uso de plataforma')

    # Métricas disponibles
    add_heading(doc, '6.10 Métricas Disponibles', 2)
    if 'metricas_informes' in data:
        metricas = data['metricas_informes']
        add_paragraph_text(doc,
            'El plugin proporciona diversas métricas organizadas en categorías para el análisis '
            'del uso de la plataforma:')

        if 'categorias' in metricas:
            categorias_data = [['Categoría', 'Métricas principales']]
            for cat in metricas['categorias'][:6]:
                nombre = cat.get('nombre', 'N/A')
                mets = cat.get('metricas', [])
                mets_str = ', '.join(mets[:3]) if mets else 'N/A'
                if len(mets_str) > 60:
                    mets_str = mets_str[:57] + '...'
                categorias_data.append([nombre, mets_str])
            add_table_with_caption(doc, categorias_data, 'Categorías de métricas del plugin report_platform_usage')

    # Algoritmo de dedicación
    add_heading(doc, '6.11 Algoritmo de Cálculo de Dedicación', 2)
    if 'algoritmo_dedicacion' in data:
        alg = data['algoritmo_dedicacion']
        add_paragraph_text(doc,
            f'{alg.get("descripcion", "Algoritmo basado en sesiones para calcular tiempo de dedicación.")}')

        if 'parametros' in alg:
            params = alg['parametros']
            params_data = [['Parámetro', 'Descripción', 'Valor por defecto']]
            for key, val in params.items():
                params_data.append([
                    key,
                    val.get('descripcion', 'N/A')[:40] + '...' if len(val.get('descripcion', '')) > 40 else val.get('descripcion', 'N/A'),
                    val.get('defecto', 'N/A')
                ])
            add_table_with_caption(doc, params_data, 'Parámetros del algoritmo de dedicación')

    # Sistema de caché
    add_heading(doc, '6.12 Sistema de Caché', 2)
    if 'sistema_cache' in data:
        cache = data['sistema_cache']
        add_paragraph_text(doc,
            f'El plugin implementa un sistema de caché con estrategia "{cache.get("estrategia", "cache-aside")}" '
            'para optimizar el rendimiento de las consultas.')

        if 'definiciones' in cache:
            cache_data = [['Nombre', 'Modo', 'TTL', 'Propósito']]
            for c in cache['definiciones']:
                cache_data.append([
                    c.get('nombre', 'N/A'),
                    c.get('modo', 'N/A'),
                    c.get('ttl', 'N/A'),
                    c.get('proposito', 'N/A')[:35] + '...' if len(c.get('proposito', '')) > 35 else c.get('proposito', 'N/A')
                ])
            add_table_with_caption(doc, cache_data, 'Definiciones de caché del plugin report_platform_usage')

    # Casos de uso
    add_heading(doc, '6.13 Casos de Uso', 2)
    if 'casos_uso' in data:
        add_paragraph_text(doc,
            'El plugin está diseñado para satisfacer los siguientes escenarios de uso:')

        casos_data = [['Caso de Uso', 'Beneficio']]
        for caso in data['casos_uso'][:4]:
            casos_data.append([
                caso.get('caso', 'N/A'),
                caso.get('beneficio', 'N/A')
            ])
        add_table_with_caption(doc, casos_data, 'Casos de uso del plugin report_platform_usage')

    # Integración con el sistema
    add_heading(doc, '6.14 Integración con Moodle e IOMAD', 2)
    if 'integracion_sistema' in data:
        integ = data['integracion_sistema']
        add_paragraph_text(doc,
            'El plugin se integra con los sistemas de navegación de Moodle y detecta '
            'automáticamente la presencia de IOMAD para habilitar el filtrado por empresas.')

        if 'integracion_iomad' in integ:
            iomad = integ['integracion_iomad']
            if 'funcionalidad' in iomad:
                add_bullet_list(doc, iomad['funcionalidad'])

    # Interfaz de usuario
    add_heading(doc, '6.15 Interfaz de Usuario', 2)
    if 'interfaz_usuario' in data:
        ui = data['interfaz_usuario']
        if 'paginas' in ui:
            add_paragraph_text(doc,
                'El plugin proporciona varias páginas web para la interacción con el usuario:')

            pages_data = [['Página', 'Propósito', 'Componentes']]
            for page in ui['paginas']:
                components = page.get('componentes', [])
                comp_str = ', '.join(components[:2]) if components else 'N/A'
                if len(comp_str) > 40:
                    comp_str = comp_str[:37] + '...'
                pages_data.append([
                    page.get('archivo', 'N/A'),
                    page.get('proposito', 'N/A')[:35] + '...' if len(page.get('proposito', '')) > 35 else page.get('proposito', 'N/A'),
                    comp_str
                ])
            add_table_with_caption(doc, pages_data, 'Páginas de interfaz del plugin report_platform_usage')

        if 'tecnologias_frontend' in ui:
            add_paragraph_text(doc, 'Tecnologías utilizadas en el frontend:')
            add_bullet_list(doc, ui['tecnologias_frontend'])

    # Configuración de administración
    add_heading(doc, '6.16 Configuración de Administración', 2)
    if 'configuracion_administracion' in data:
        config = data['configuracion_administracion']
        add_paragraph_text(doc,
            f'El plugin se configura desde: {config.get("ubicacion_menu", "Administración del sitio")}')

        if 'parametros_configurables' in config:
            params_data = [['Parámetro', 'Tipo', 'Defecto', 'Impacto']]
            for param in config['parametros_configurables']:
                params_data.append([
                    param.get('nombre', 'N/A'),
                    param.get('tipo', 'N/A'),
                    str(param.get('defecto', 'N/A')),
                    param.get('impacto', 'N/A')
                ])
            add_table_with_caption(doc, params_data, 'Parámetros configurables del plugin report_platform_usage')

    # Rendimiento y escalabilidad
    add_heading(doc, '6.17 Rendimiento y Escalabilidad', 2)
    if 'rendimiento_escalabilidad' in data:
        rend = data['rendimiento_escalabilidad']
        add_paragraph_text(doc,
            'El plugin implementa diversas optimizaciones para garantizar un buen rendimiento:')

        if 'optimizaciones' in rend:
            add_bullet_list(doc, rend['optimizaciones'][:6])

        if 'tiempos_cache' in rend:
            tc = rend['tiempos_cache']
            cache_info = [
                f'Datos del informe: {tc.get("datos_informe", "N/A")}',
                f'Listas de usuarios: {tc.get("listas_usuarios", "N/A")}',
                f'Razón: {tc.get("razon", "N/A")}'
            ]
            add_bullet_list(doc, cache_info)

    # Internacionalización
    add_heading(doc, '6.18 Internacionalización', 2)
    if 'internacionalizacion' in data:
        i18n = data['internacionalizacion']
        add_paragraph_text(doc,
            f'El plugin soporta {len(i18n.get("idiomas_soportados", []))} idiomas: '
            f'{", ".join(i18n.get("idiomas_soportados", ["N/A"]))}. '
            f'Contiene aproximadamente {i18n.get("total_cadenas", "N/A")} cadenas de texto '
            f'con facilidad de traducción {i18n.get("facilidad_traduccion", "alta")}.')

    # Limitaciones conocidas
    add_heading(doc, '6.19 Limitaciones Conocidas', 2)
    if 'limitaciones_conocidas' in data:
        add_paragraph_text(doc,
            'El plugin tiene las siguientes limitaciones que deben considerarse:')
        add_bullet_list(doc, data['limitaciones_conocidas'][:6])

    # Mejores prácticas
    add_heading(doc, '6.20 Mejores Prácticas', 2)
    if 'mejores_practicas' in data:
        mp = data['mejores_practicas']
        if 'configuracion' in mp:
            add_paragraph_text(doc, 'Configuración recomendada:', bold=True)
            add_bullet_list(doc, mp['configuracion'][:4])
        if 'uso' in mp:
            add_paragraph_text(doc, 'Uso óptimo del plugin:', bold=True)
            add_bullet_list(doc, mp['uso'][:4])
        if 'rendimiento' in mp:
            add_paragraph_text(doc, 'Optimización del rendimiento:', bold=True)
            add_bullet_list(doc, mp['rendimiento'][:4])

    # Conclusiones del plugin
    add_heading(doc, '6.21 Conclusiones del Plugin', 2)
    if 'conclusiones' in data:
        conc = data['conclusiones']
        if 'fortalezas' in conc:
            add_paragraph_text(doc, 'Fortalezas principales:', bold=True)
            add_bullet_list(doc, conc['fortalezas'][:5])
        if 'areas_mejora' in conc:
            add_paragraph_text(doc, 'Áreas de mejora identificadas:', bold=True)
            add_bullet_list(doc, conc['areas_mejora'][:4])
        if 'recomendacion_uso' in conc:
            add_paragraph_text(doc, f'Recomendación: {conc["recomendacion_uso"]}')

    doc.add_page_break()


def generate_platform_access_section(doc):
    """Genera la sección del plugin local_platform_access."""
    data = load_json_data('local_platform_access')
    if not data:
        return

    add_heading(doc, '7. PLUGIN LOCAL_PLATFORM_ACCESS', 1)

    # Introducción
    add_heading(doc, '7.1 Descripción General', 2)
    if 'resumen_ejecutivo' in data:
        resumen = data['resumen_ejecutivo']
        add_paragraph_text(doc,
            f'{data.get("nombre_completo", "Local Platform Access")} es un plugin de tipo '
            f'{data.get("tipo", "local")} cuyo propósito es '
            f'{resumen.get("proposito", "generar datos de acceso para pruebas")}.')
    else:
        add_paragraph_text(doc,
            'Este plugin permite generar registros de acceso simulados para realizar pruebas '
            'de carga y validar el rendimiento del sistema de reportes.')

    # Estructura
    add_heading(doc, '7.2 Estructura del Plugin', 2)
    svg_path = SVG_DIR / 'local_platform_access' / 'estructura_directorios.svg'
    add_figure(doc, svg_path, 'Estructura de directorios del plugin local_platform_access')

    # Información de versión
    add_heading(doc, '7.3 Información de Versión', 2)
    if 'version' in data:
        v = data['version']
        version_data = [
            ['Atributo', 'Valor'],
            ['Plugin', data.get('plugin', 'local_platform_access')],
            ['Versión', v.get('release', 'N/A')],
            ['Moodle requerido', v.get('requires', 'N/A')],
            ['Madurez', v.get('maturity', 'N/A')],
            ['Licencia', v.get('licencia', 'GPLv3')]
        ]
        add_table_with_caption(doc, version_data, 'Información de versión del plugin local_platform_access')

    # Arquitectura
    add_heading(doc, '7.4 Arquitectura del Sistema', 2)
    if 'arquitectura' in data:
        arq = data['arquitectura']
        add_paragraph_text(doc,
            f'El plugin implementa el patrón de diseño {arq.get("patron_diseno", "Singleton")}, '
            'optimizado para operaciones masivas de inserción en base de datos y generación '
            'eficiente de datos aleatorios.')

    svg_path = SVG_DIR / 'local_platform_access' / 'arquitectura.svg'
    add_figure(doc, svg_path, 'Diagrama de arquitectura del plugin local_platform_access')

    # Capabilities
    add_heading(doc, '7.5 Capabilities y Seguridad', 2)
    add_paragraph_text(doc,
        'Dado el potencial impacto en el rendimiento de la base de datos, el plugin '
        'implementa controles estrictos de acceso.')

    if 'capacidades' in data and 'lista' in data['capacidades']:
        caps = data['capacidades']['lista']
        caps_data = [['Capability', 'Tipo', 'Propósito']]
        for cap in caps:
            caps_data.append([
                cap.get('nombre', 'N/A'),
                cap.get('tipo', 'N/A'),
                cap.get('proposito', 'N/A')
            ])
        add_table_with_caption(doc, caps_data, 'Capabilities del plugin local_platform_access')

    # Flujo de generación
    add_heading(doc, '7.6 Flujo de Generación de Datos', 2)
    add_paragraph_text(doc,
        'El proceso de generación permite configurar la cantidad de registros, '
        'el rango de fechas y los usuarios afectados, optimizando las inserciones '
        'mediante procesamiento por lotes.')

    svg_path = SVG_DIR / 'local_platform_access' / 'diagrama_generacion.svg'
    add_figure(doc, svg_path, 'Diagrama de flujo de generación del plugin local_platform_access')

    # Flujo de acceso
    add_heading(doc, '7.7 Flujo de Acceso al Plugin', 2)
    add_paragraph_text(doc,
        'El acceso al plugin sigue un flujo controlado que incluye validación de permisos, '
        'confirmación del usuario y ejecución supervisada con barra de progreso.')

    svg_path = SVG_DIR / 'local_platform_access' / 'flujo_acceso.svg'
    add_figure(doc, svg_path, 'Flujo de acceso y ejecución del plugin local_platform_access')

    # Optimizaciones
    add_heading(doc, '7.8 Optimizaciones de Rendimiento', 2)
    if 'optimizaciones_rendimiento' in data and 'tecnicas' in data['optimizaciones_rendimiento']:
        opt = data['optimizaciones_rendimiento']
        add_paragraph_text(doc,
            'Para manejar la generación masiva de registros, el plugin implementa '
            'varias técnicas de optimización:')

        opt_data = [['Técnica', 'Descripción', 'Impacto']]
        for t in opt['tecnicas']:
            opt_data.append([
                t.get('nombre', 'N/A'),
                t.get('descripcion', 'N/A')[:40] + '...' if len(t.get('descripcion', '')) > 40 else t.get('descripcion', 'N/A'),
                t.get('impacto', 'N/A')
            ])
        add_table_with_caption(doc, opt_data, 'Técnicas de optimización implementadas')

    # Base de datos utilizada
    add_heading(doc, '7.9 Base de Datos', 2)
    add_paragraph_text(doc,
        'Este plugin no crea tablas propias, pero utiliza múltiples tablas existentes de Moodle '
        'e IOMAD para generar e insertar los registros de acceso.')

    if 'base_datos' in data:
        bd = data['base_datos']
        if 'tablas_estandar_utilizadas' in bd:
            tables_data = [['Tabla', 'Operación', 'Propósito']]
            for t in bd['tablas_estandar_utilizadas']:
                tables_data.append([
                    t.get('tabla', 'N/A'),
                    t.get('operacion', 'N/A'),
                    t.get('proposito', 'N/A')[:50] + '...' if len(t.get('proposito', '')) > 50 else t.get('proposito', 'N/A')
                ])
            add_table_with_caption(doc, tables_data, 'Tablas de Moodle utilizadas por local_platform_access')

    # Casos de uso
    add_heading(doc, '7.10 Casos de Uso', 2)
    if 'casos_uso' in data:
        add_paragraph_text(doc,
            'El plugin está diseñado para los siguientes escenarios:')
        add_bullet_list(doc, data['casos_uso'])

    # Integración IOMAD
    add_heading(doc, '7.11 Integración con IOMAD', 2)
    if 'integracion_iomad' in data:
        iomad = data['integracion_iomad']
        add_paragraph_text(doc,
            'El plugin se integra completamente con IOMAD para entornos multi-empresa, '
            'permitiendo filtrar y generar datos por empresa específica.')

        if 'funcionalidades' in iomad:
            add_bullet_list(doc, iomad['funcionalidades'])

        add_paragraph_text(doc,
            f'{iomad.get("compatibilidad", "Compatible con Moodle estándar.")}')

    # Páginas de interfaz
    add_heading(doc, '7.12 Páginas de Interfaz', 2)
    if 'paginas_interfaz' in data:
        add_paragraph_text(doc,
            'El plugin proporciona las siguientes páginas de interfaz:')

        pages_data = [['Archivo', 'URL', 'Propósito']]
        for page in data['paginas_interfaz']:
            pages_data.append([
                page.get('archivo', 'N/A'),
                page.get('url', 'N/A'),
                page.get('proposito', 'N/A')
            ])
        add_table_with_caption(doc, pages_data, 'Páginas de interfaz del plugin local_platform_access')

    # Funcionalidades principales
    add_heading(doc, '7.13 Funcionalidades Principales', 2)
    if 'funcionalidades_principales' in data:
        func = data['funcionalidades_principales']

        if 'generacion_accesos' in func:
            gen = func['generacion_accesos']
            add_paragraph_text(doc, 'Generación de accesos:', bold=True)
            if 'tipos_soportados' in gen:
                add_bullet_list(doc, gen['tipos_soportados'])
            if 'eventos_generados' in gen:
                add_paragraph_text(doc, 'Eventos generados:')
                add_bullet_list(doc, gen['eventos_generados'][:6])

        if 'eventos_avanzados' in func:
            ev = func['eventos_avanzados']
            add_paragraph_text(doc, 'Eventos avanzados:', bold=True)
            ev_list = []
            for key, val in ev.items():
                ev_list.append(f'{key.title()}: {val}')
            add_bullet_list(doc, ev_list)

    # Estadísticas generadas
    add_heading(doc, '7.14 Estadísticas Generadas', 2)
    if 'estadisticas_generadas' in data:
        add_paragraph_text(doc,
            'El plugin proporciona estadísticas detalladas después de cada ejecución:')
        add_bullet_list(doc, data['estadisticas_generadas'][:10])

    # Seguridad y privacidad
    add_heading(doc, '7.15 Seguridad y Privacidad', 2)
    if 'seguridad_privacidad' in data:
        seg = data['seguridad_privacidad']
        add_paragraph_text(doc,
            f'GDPR Compliance: {"Sí" if seg.get("gdpr_compliance") else "No"}. '
            f'Almacena datos personales: {"Sí" if seg.get("almacena_datos_personales") else "No"}.')

        if 'controles_seguridad' in seg:
            add_paragraph_text(doc, 'Controles de seguridad implementados:')
            add_bullet_list(doc, seg['controles_seguridad'])

        if 'capacidades_riesgo' in seg:
            add_paragraph_text(doc, 'Riesgos asociados a las capacidades:')
            add_bullet_list(doc, seg['capacidades_riesgo'])

    # Limitaciones y consideraciones
    add_heading(doc, '7.16 Limitaciones y Consideraciones', 2)
    if 'limitaciones_consideraciones' in data:
        add_paragraph_text(doc,
            'Es importante considerar las siguientes limitaciones del plugin:')
        add_bullet_list(doc, data['limitaciones_consideraciones'])

    # Métricas de código
    add_heading(doc, '7.17 Métricas de Código', 2)
    if 'metricas_codigo' in data:
        mc = data['metricas_codigo']
        metrics_data = [['Métrica', 'Valor']]
        for key, val in mc.items():
            metrics_data.append([key.replace('_', ' ').title(), str(val)])
        add_table_with_caption(doc, metrics_data, 'Métricas de código del plugin local_platform_access')

    # Flujo completo
    add_heading(doc, '7.18 Flujo Completo de Ejecución', 2)
    if 'flujo_completo' in data:
        fc = data['flujo_completo']
        add_paragraph_text(doc,
            'El siguiente flujo describe el proceso completo de generación de datos:')
        flujo_list = []
        for key, val in sorted(fc.items()):
            flujo_list.append(f'{key.replace("paso_", "Paso ")}: {val}')
        add_bullet_list(doc, flujo_list[:12])

    # Dependencias
    add_heading(doc, '7.19 Dependencias', 2)
    if 'dependencias' in data:
        dep = data['dependencias']
        add_paragraph_text(doc,
            f'Versión mínima de Moodle requerida: {dep.get("moodle_version", "N/A")}')

        if 'plugins_opcionales' in dep:
            add_paragraph_text(doc, 'Plugins opcionales para funcionalidad extendida:')
            add_bullet_list(doc, dep['plugins_opcionales'])

    # Conclusión del plugin
    add_heading(doc, '7.20 Conclusión', 2)
    if 'conclusion' in data:
        conc = data['conclusion']
        conc_list = []
        for key, val in conc.items():
            conc_list.append(f'{key.replace("_", " ").title()}: {val}')
        add_bullet_list(doc, conc_list)

    doc.add_page_break()


def generate_conclusions(doc):
    """Genera la sección de conclusiones."""
    add_heading(doc, '8. CONCLUSIONES', 1)

    add_paragraph_text(doc,
        'Los tres plugins documentados en este informe representan extensiones significativas '
        'de la plataforma Moodle ISER, cada uno cumpliendo un propósito específico dentro del '
        'ecosistema de gestión académica de la institución.')

    conclusions = [
        'El plugin local_jobboard proporciona una solución completa para la gestión del proceso '
        'de contratación docente, desde la publicación de vacantes hasta la formalización de contratos.',
        'El plugin report_platform_usage permite monitorear el uso efectivo de la plataforma, '
        'proporcionando métricas valiosas para la toma de decisiones institucionales.',
        'El plugin local_platform_access facilita las pruebas de carga y rendimiento mediante '
        'la generación controlada de datos de acceso simulados.'
    ]
    add_bullet_list(doc, conclusions)

    add_heading(doc, '9. RECOMENDACIONES', 1)

    recommendations = [
        'Mantener actualizada la documentación técnica ante cualquier cambio en los plugins.',
        'Realizar pruebas de regresión antes de cada actualización de Moodle.',
        'Implementar pruebas automatizadas para los servicios web críticos.',
        'Monitorear el rendimiento de las consultas a base de datos en producción.',
        'Establecer un proceso de revisión de código para cualquier modificación.'
    ]
    add_bullet_list(doc, recommendations)


def generate_back_cover(doc):
    """Genera la contraportada institucional del documento."""
    doc.add_page_break()

    # Espacio inicial
    for _ in range(4):
        doc.add_paragraph()

    # Encabezado institucional
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run('INSTITUTO SUPERIOR DE EDUCACIÓN RURAL')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    run.font.color.rgb = COLORS['VERDE']

    header2 = doc.add_paragraph()
    header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header2.add_run('ISER')
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = 'Arial'
    run.font.color.rgb = COLORS['VERDE']

    # Línea decorativa
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run('═' * 50)
    run.font.size = Pt(12)
    run.font.color.rgb = COLORS['AMARILLO']

    for _ in range(3):
        doc.add_paragraph()

    # Título de sección
    info_title = doc.add_paragraph()
    info_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_title.add_run('INFORMACIÓN DEL CONTRATISTA')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    run.font.color.rgb = COLORS['GRIS']

    for _ in range(2):
        doc.add_paragraph()

    # Nombre
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run('PEDRO ALONSO ARIAS BALCUCHO')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Arial'

    # CC
    cc = doc.add_paragraph()
    cc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cc.add_run('C.C. 1.094.277.917')
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.font.color.rgb = COLORS['GRIS']

    doc.add_paragraph()

    # Contrato
    contract = doc.add_paragraph()
    contract.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contract.add_run('Contrato de Prestación de Servicios')
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    run.font.color.rgb = COLORS['GRIS']

    contract_num = doc.add_paragraph()
    contract_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contract_num.add_run('No. CD-CPS-066-2025')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    run.font.color.rgb = COLORS['VERDE']

    for _ in range(3):
        doc.add_paragraph()

    # Recuadro de contacto
    box_line = doc.add_paragraph()
    box_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = box_line.add_run('┌' + '─' * 36 + '┐')
    run.font.size = Pt(10)
    run.font.name = 'Consolas'
    run.font.color.rgb = COLORS['VERDE']

    contact_header = doc.add_paragraph()
    contact_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact_header.add_run('│          DATOS DE CONTACTO          │')
    run.font.size = Pt(10)
    run.font.name = 'Consolas'
    run.font.color.rgb = COLORS['VERDE']

    phone = doc.add_paragraph()
    phone.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = phone.add_run('│     Tel: 3112243871                 │')
    run.font.size = Pt(10)
    run.font.name = 'Consolas'
    run.font.color.rgb = COLORS['VERDE']

    email = doc.add_paragraph()
    email.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = email.add_run('│     alonsoariasb@gmail.com          │')
    run.font.size = Pt(10)
    run.font.name = 'Consolas'
    run.font.color.rgb = COLORS['VERDE']

    box_end = doc.add_paragraph()
    box_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = box_end.add_run('└' + '─' * 36 + '┘')
    run.font.size = Pt(10)
    run.font.name = 'Consolas'
    run.font.color.rgb = COLORS['VERDE']

    for _ in range(4):
        doc.add_paragraph()

    # Línea decorativa inferior
    line2 = doc.add_paragraph()
    line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line2.add_run('═' * 50)
    run.font.size = Pt(12)
    run.font.color.rgb = COLORS['AMARILLO']

    # Lugar y fecha
    location = doc.add_paragraph()
    location.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = location.add_run('Pamplona, Norte de Santander - Colombia')
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    run.font.color.rgb = COLORS['GRIS']

    year = doc.add_paragraph()
    year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = year.add_run('Diciembre 2025')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.font.color.rgb = COLORS['VERDE']

    doc.add_page_break()


def main():
    """Función principal."""
    print('=' * 60)
    print('  GENERADOR DE INFORME TÉCNICO - FORMATO F-GCT-17')
    print('  (Versión Python - Edición de documento existente)')
    print('=' * 60)
    print()

    # Verificar que existe el template
    if not TEMPLATE_PATH.exists():
        print(f'❌ Error: No se encontró el template: {TEMPLATE_PATH}')
        return 1

    print(f'📄 Abriendo template: {TEMPLATE_PATH.name}')

    # Abrir documento template
    doc = Document(str(TEMPLATE_PATH))

    # Generar contenido
    print('  📝 Generando portada...')
    generate_cover_page(doc)

    print('  📄 Generando contraportada...')
    generate_back_cover(doc)

    print('  📑 Generando tablas de contenido...')
    generate_toc_pages(doc)

    print('  📝 Generando secciones introductorias...')
    generate_introduction(doc)

    print('  🔧 Generando sección local_jobboard...')
    generate_jobboard_section(doc)

    print('  📊 Generando sección report_platform_usage...')
    generate_platform_usage_section(doc)

    print('  🔐 Generando sección local_platform_access...')
    generate_platform_access_section(doc)

    print('  📋 Generando conclusiones...')
    generate_conclusions(doc)

    # Guardar documento
    print(f'  💾 Guardando documento: {OUTPUT_PATH.name}')
    doc.save(str(OUTPUT_PATH))

    print()
    print('=' * 60)
    print('✅ Documento generado exitosamente:')
    print(f'   {OUTPUT_PATH}')
    print()
    print(f'📊 Estadísticas:')
    print(f'   - Figuras generadas: {figura_counter}')
    print(f'   - Tablas generadas: {tabla_counter}')
    print()
    print('📋 Instrucciones post-generación:')
    print('   1. Abra el documento en Microsoft Word')
    print('   2. Presione Ctrl+A para seleccionar todo')
    print('   3. Presione F9 para actualizar los campos')
    print('   4. Verifique las tablas de contenido')
    print('=' * 60)

    return 0


if __name__ == '__main__':
    exit(main())
